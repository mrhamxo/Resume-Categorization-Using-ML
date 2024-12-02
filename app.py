import os
import pandas as pd
import pickle
from pypdf import PdfReader
import re
import streamlit as st
from docx import Document

# Load models
word_vector = pickle.load(open("models/tfidf.pkl", "rb"))
model = pickle.load(open("models/model.pkl", "rb"))

# List of skills for extracting from resumes (can be expanded further)
skills_keywords = ['python', 'java', 'c++', 'sql', 'machine learning', 'data science', 'tensorflow', 'keras', 'html', 'css', 'javaScript', 'django', 'aws', 'docker', 'hadoop', 'spark', 'database', 'etl']

def cleanResume(txt):
    cleanText = re.sub(r'http\S+\s', ' ', txt)
    cleanText = re.sub(r'RT|cc', ' ', cleanText)
    cleanText = re.sub(r'#\S+\s', ' ', cleanText)
    cleanText = re.sub(r'@\S+', '  ', cleanText)
    cleanText = re.sub(r'[{}]'.format(re.escape("""!"#$%&'()*+,-./:;<=>?@[/]^_`{|}~""")), ' ', cleanText)
    cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
    cleanText = re.sub(r'\s+', ' ', cleanText)
    return cleanText

category_mapping = {
    15: "Java Developer",
    23: "Testing",
    8: "DevOps Engineer",
    20: "Python Developer",
    24: "Web Designing",
    12: "HR",
    13: "Hadoop",
    3: "Blockchain",
    10: "ETL Developer",
    18: "Operations Manager",
    6: "Data Science",
    22: "Sales",
    16: "Mechanical Engineer",
    1: "Arts",
    7: "Database",
    11: "Electrical Engineering",
    14: "Health and fitness",
    19: "PMO",
    4: "Business Analyst",
    9: "DotNet Developer",
    2: "Automation Testing",
    17: "Network Security Engineer",
    21: "SAP Developer",
    5: "Civil Engineer",
    0: "Advocate",
}

def extract_text_from_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def extract_skills(text):
    skills = [skill for skill in skills_keywords if skill.lower() in text.lower()]
    return ", ".join(skills) if skills else "Not Mentioned"

def extract_experience(text):
    # Simple regex to find years of experience (e.g., "5 years of experience")
    experience_match = re.search(r'(\d+)\s*year[s]?\s*experience', text, re.IGNORECASE)
    return experience_match.group(1) if experience_match else "Not Mentioned"

def extract_location(text):
    # Simple location extraction (you can refine with more sophisticated NER or regex)
    location_keywords = ['new york', 'california', 'texas', 'london', 'mumbai', 'delhi', 'los angeles']
    for location in location_keywords:
        if location.lower() in text.lower():
            return location.capitalize()
    return "Not Mentioned"

def categorize_resumes(uploaded_files, output_directory):
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    results = []
    
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith('.pdf'):
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            text = extract_text_from_docx(uploaded_file)
        else:
            continue
        
        cleaned_resume = cleanResume(text)

        input_features = word_vector.transform([cleaned_resume])
        prediction_id = model.predict(input_features)[0]
        category_name = category_mapping.get(prediction_id, "Unknown")
        
        skills = extract_skills(text)
        experience = extract_experience(text)
        location = extract_location(text)
        
        category_folder = os.path.join(output_directory, category_name)
        
        if not os.path.exists(category_folder):
            os.makedirs(category_folder)
        
        target_path = os.path.join(category_folder, uploaded_file.name)
        with open(target_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        results.append({
            'filename': uploaded_file.name,
            'category': category_name,
            'skills': skills,
            'experience': experience,
            'location': location
        })
    
    results_df = pd.DataFrame(results)
    return results_df

def generate_docx(results_df, docx_filename):
    # Create a new DOCX document
    doc = Document()
    doc.add_heading('Resume Categorization Results', 0)
    
    # Add the results to the DOCX file
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filename'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Skills'
    hdr_cells[3].text = 'Experience'
    hdr_cells[4].text = 'Location'
    
    # Add the rows of results
    for index, row in results_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['filename']
        row_cells[1].text = row['category']
        row_cells[2].text = row['skills']
        row_cells[3].text = row['experience']
        row_cells[4].text = row['location']
    
    # Save the DOCX file
    doc.save(docx_filename)

st.title("Advanced Resume Categorizer Application")
st.subheader("With Python & Machine Learning")

uploaded_files = st.file_uploader("Choose PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)
output_directory = st.text_input("Output Directory", "categorized_resumes")

if st.button("Categorize Resumes"):
    if uploaded_files and output_directory:
        results_df = categorize_resumes(uploaded_files, output_directory)
        st.write(results_df)
        
        # Save results as CSV
        results_csv = results_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Download results as CSV",
            data=results_csv,
            file_name='categorized_resumes.csv',
            mime='text/csv',
        )
        
        # Save results as DOCX
        docx_filename = os.path.join(output_directory, 'categorized_resumes.docx')
        generate_docx(results_df, docx_filename)
        with open(docx_filename, 'rb') as docx_file:
            st.download_button(
                label="Download results as DOCX",
                data=docx_file,
                file_name='categorized_resumes.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        
        st.success("Resumes categorization and processing completed.")
    else:
        st.error("Please upload files and specify the output directory.")
